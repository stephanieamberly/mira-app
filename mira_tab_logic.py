import streamlit as st
import openai
from datetime import datetime, timedelta
import sqlite3
import os
import base64
import dateparser
import re
import csv
from io import StringIO
import pdfplumber
from docx import Document
from docx.shared import Pt

DB_FILE = "mira_resumes.db"

# --- HELPER FUNCTIONS ---
def ask_gpt(prompt):
    openai.api_key = st.secrets["openai"]["api_key"]
    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are MIRA, an intelligent, helpful, and friendly AI recruiting assistant."},
            {"role": "user", "content": prompt}
        ]
    )
    return response.choices[0].message.content.strip()

def extract_text_from_pdf(file):
    with pdfplumber.open(file) as pdf:
        return "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())

def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_details(text):
    name = text.split("\n")[0].strip() if text else ""
    email_match = re.search(r"[\w\.-]+@[\w\.-]+", text)
    phone_match = re.search(r"(\+\d{1,2}\s)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text)
    email = email_match.group() if email_match else ""
    phone = phone_match.group() if phone_match else ""
    skills = experience = ""
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if "skills" in line.lower():
            skills = "\n".join(lines[i+1:i+6])
        if "experience" in line.lower():
            experience = "\n".join(lines[i+1:i+10])
    return name, email, phone, skills, experience

def save_to_db(name, email, phone, skills, experience, filename):
    conn = sqlite3.connect(DB_FILE)
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO resumes (name, email, phone, skills, experience, filename, timestamp)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (name, email, phone, skills, experience, filename, datetime.now().isoformat()))
    conn.commit()
    conn.close()

def schedule_google_event(candidate_name, candidate_email, interview_date, interview_time, position_title, teams_link="https://teams.microsoft.com/l/meetup-join/abc123"):
    return teams_link or "https://teams.microsoft.com/l/meetup-join/abc123"

def generate_onboarding_doc(name, email, position, start_date, salary):
    doc = Document()
    doc.add_heading('Offer Letter', 0)
    doc.add_paragraph(f"Dear {name},")
    doc.add_paragraph(f"We are excited to offer you the position of {position}. Your start date will be {start_date}, with a starting salary of ${salary}.")
    doc.add_paragraph("Please let us know if you have any questions.")
    doc.add_paragraph("Sincerely,\nHR Team")

    for para in doc.paragraphs:
        for run in para.runs:
            run.font.size = Pt(11)

    filepath = f"onboarding_docs/{name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    os.makedirs("onboarding_docs", exist_ok=True)
    doc.save(filepath)

    conn = sqlite3.connect(DB_FILE)
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO onboarding_logs (name, email, position, start_date, salary, filepath, timestamp)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (name, email, position, start_date, salary, filepath, datetime.now().isoformat()))
    conn.commit()
    conn.close()

    return filepath

# --- DB SETUP ---
def init_db():
    conn = sqlite3.connect(DB_FILE)
    cur = conn.cursor()

    cur.execute("""
    CREATE TABLE IF NOT EXISTS resumes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT,
        email TEXT,
        phone TEXT,
        skills TEXT,
        experience TEXT,
        filename TEXT,
        timestamp TEXT,
        job_title TEXT DEFAULT '',
        status TEXT DEFAULT 'New',
        score INTEGER DEFAULT 0
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS mira_logs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        question TEXT,
        answer TEXT,
        timestamp TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS onboarding_logs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT,
        email TEXT,
        position TEXT,
        start_date TEXT,
        salary TEXT,
        filepath TEXT,
        timestamp TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS job_descriptions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        content TEXT,
        timestamp TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS branding_assets (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT,
        content TEXT,
        timestamp TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS feedback_surveys (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        candidate_name TEXT,
        rating INTEGER,
        comments TEXT,
        timestamp TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS coaching_materials (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        title TEXT,
        notes TEXT,
        timestamp TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS analytics_snapshots (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        metric TEXT,
        value INTEGER,
        timestamp TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS voice_assistant_logs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        transcript TEXT,
        timestamp TEXT
    )
    """)

    conn.commit()
    conn.close()

# --- RENDER TABS ---
def render_tabs(tab1, tab2, tab3, tab4, tab5, tab6, tab7, tab8):
    with tab1:
        st.subheader("🧠 Ask MIRA")

        if st.button("🎤 Start Listening"):
            st.info("Voice input activated (simulated)")

        user_input = st.text_input("Ask me anything related to recruiting, HR, or employer branding:")
        if user_input:
            response = ask_gpt(user_input)
            st.markdown(f"**MIRA says:** {response}")

    with tab2:
        st.subheader("📄 Resume Viewer")

        filter = st.text_input("Search resumes by name, email, skills, or experience")

        conn = sqlite3.connect(DB_FILE)
        cur = conn.cursor()
        if filter:
            cur.execute("""
                SELECT * FROM resumes
                WHERE name LIKE ? OR email LIKE ? OR skills LIKE ? OR experience LIKE ?
            """, (f"%{filter}%",)*4)
        else:
            cur.execute("SELECT * FROM resumes")
        rows = cur.fetchall()
        conn.close()

        for r in rows:
            st.markdown(f"**{r[1]}** | {r[2]} | {r[3]}")
            st.markdown(f"**Skills:** {r[4][:150]}...")
            st.markdown(f"**Experience:** {r[5][:200]}...")
            st.markdown("---")

        st.subheader("📤 Upload Resume")
        uploaded_file = st.file_uploader("Upload a resume (.pdf or .docx)", type=["pdf", "docx"])
        if uploaded_file:
            ext = uploaded_file.name.split(".")[-1].lower()
            raw_text = extract_text_from_pdf(uploaded_file) if ext == "pdf" else extract_text_from_docx(uploaded_file)
            name, email, phone, skills, experience = extract_details(raw_text)
            save_to_db(name, email, phone, skills, experience, uploaded_file.name)
            st.success(f"Saved resume for: {name}")

    with tab3:
        st.subheader("📅 Calendar & Interview Scheduling")

        with st.form("calendar_form"):
            candidate_name = st.text_input("Candidate Name")
            candidate_email = st.text_input("Candidate Email")
            position_title = st.text_input("Position Title")
            interview_date = st.date_input("Interview Date")
            interview_time = st.time_input("Interview Time")
            teams_link = st.text_input("Microsoft Teams Link (Paste here)")
            submitted = st.form_submit_button("📅 Schedule Interview")

            if submitted:
                try:
                    link = schedule_google_event(
                        candidate_name,
                        candidate_email,
                        interview_date.strftime("%Y-%m-%d"),
                        interview_time.strftime("%H:%M"),
                        position_title,
                        teams_link
                    )
                    st.success(f"Interview scheduled! Join via [Teams]({link})")
                except Exception as e:
                    st.error(f"Error: {e}")

    with tab4:
        st.subheader("📁 Onboarding Documents")

        with st.form("onboarding_form"):
            name = st.text_input("Candidate Name")
            email = st.text_input("Candidate Email")
            position = st.text_input("Position Title")
            start_date = st.date_input("Start Date")
            salary = st.number_input("Salary Offered", min_value=30000, step=500)
            submitted = st.form_submit_button("📄 Generate Offer Letter")

            if submitted:
                filepath = generate_onboarding_doc(name, email, position, start_date, salary)
                st.success(f"Offer letter generated for {name}")

                if os.path.exists(filepath):
                    with open(filepath, "rb") as f:
                        file_data = f.read()
                        st.download_button(
                            "⬇️ Download Document",
                            data=file_data,
                            file_name=os.path.basename(filepath),
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                else:
                    st.error("File could not be found. Please try again.")

        # Show existing generated docs
        conn = sqlite3.connect(DB_FILE)
        docs = conn.execute("SELECT name, email, position, start_date, salary, filepath, timestamp FROM onboarding_logs ORDER BY timestamp DESC").fetchall()
        conn.close()

        for d in docs:
            st.markdown(f"**{d[0]}** | {d[1]} | {d[2]} | Start: {d[3]} | 💰 ${d[4]}")
            st.markdown(f"[📄 Download]({d[5]}) — Created: *{d[6]}*")
            st.markdown("---")

    with tab5:
        st.subheader("📂 Job Description Hub")

        with st.form("jd_form"):
            jd_content = st.text_area("Paste or write a job description")
            submitted = st.form_submit_button("💾 Save JD")
            if submitted and jd_content.strip():
                conn = sqlite3.connect(DB_FILE)
                conn.execute("INSERT INTO job_descriptions (content, timestamp) VALUES (?, ?)", (jd_content, datetime.now().isoformat()))
                conn.commit()
                conn.close()
                st.success("Job description saved!")

        st.markdown("### 📜 Saved Descriptions")
        conn = sqlite3.connect(DB_FILE)
        jds = conn.execute("SELECT content, timestamp FROM job_descriptions ORDER BY timestamp DESC").fetchall()
        conn.close()

        for content, ts in jds:
            st.code(content)
            st.caption(f"🕒 {ts}")
            st.markdown("---")

    with tab6:
        st.subheader("🎨 Employer Branding Assets")

        with st.form("branding_form"):
            name = st.text_input("Asset Name")
            content = st.text_area("Content, link, or description")
            submitted = st.form_submit_button("📥 Upload")
            if submitted and name and content:
                conn = sqlite3.connect(DB_FILE)
                conn.execute("INSERT INTO branding_assets (name, content, timestamp) VALUES (?, ?, ?)", (name, content, datetime.now().isoformat()))
                conn.commit()
                conn.close()
                st.success(f"Uploaded asset: {name}")

        conn = sqlite3.connect(DB_FILE)
        assets = conn.execute("SELECT name, content, timestamp FROM branding_assets ORDER BY timestamp DESC").fetchall()
        conn.close()

        for name, content, ts in assets:
            st.markdown(f"**{name}**")
            st.markdown(content)
            st.caption(f"🕒 {ts}")
            st.markdown("---")

    with tab7:
        st.subheader("📊 Feedback & Candidate Experience")

        with st.form("feedback_form"):
            name = st.text_input("Candidate Name")
            rating = st.slider("How was your experience?", 1, 10)
            comments = st.text_area("Additional feedback")
            submitted = st.form_submit_button("📝 Submit Feedback")
            if submitted:
                conn = sqlite3.connect(DB_FILE)
                conn.execute("INSERT INTO feedback_surveys (candidate_name, rating, comments, timestamp) VALUES (?, ?, ?, ?)", (name, rating, comments, datetime.now().isoformat()))
                conn.commit()
                conn.close()
                st.success("Thanks for your feedback!")

        st.markdown("### Recent Feedback")
        conn = sqlite3.connect(DB_FILE)
        feedback = conn.execute("SELECT candidate_name, rating, comments, timestamp FROM feedback_surveys ORDER BY timestamp DESC").fetchall()
        conn.close()

        for name, rating, comments, ts in feedback:
            st.markdown(f"**{name}** rated {rating}/10")
            st.markdown(f"💬 {comments}")
            st.caption(f"🕒 {ts}")
            st.markdown("---")

    with tab8:
        st.subheader("📈 Upskilling & Coaching")

        with st.form("coaching_form"):
            title = st.text_input("Coaching Topic")
            notes = st.text_area("Training content or notes")
            submitted = st.form_submit_button("📚 Save Entry")
            if submitted:
                conn = sqlite3.connect(DB_FILE)
                conn.execute("INSERT INTO coaching_materials (title, notes, timestamp) VALUES (?, ?, ?)", (title, notes, datetime.now().isoformat()))
                conn.commit()
                conn.close()
                st.success(f"Coaching material saved: {title}")

        conn = sqlite3.connect(DB_FILE)
        coaching = conn.execute("SELECT title, notes, timestamp FROM coaching_materials ORDER BY timestamp DESC").fetchall()
        conn.close()

        for title, notes, ts in coaching:
            st.markdown(f"### {title}")
            st.markdown(notes)
            st.caption(f"🕒 {ts}")
            st.markdown("---")