import streamlit as st
from openai import OpenAI
from datetime import datetime, timedelta
import sqlite3
import os
import base64
import dateparser
import re
import csv
from io import StringIO
import pdfplumber
from docx import Document
from docx.shared import Pt

DB_FILE = "mira_resumes.db"

# --- HELPER FUNCTIONS ---
def ask_gpt(prompt):
    client = OpenAI(api_key=st.secrets["openai"]["api_key"])
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are MIRA, an intelligent, helpful, and friendly AI recruiting assistant."},
            {"role": "user", "content": prompt}
        ]
    )
    return response.choices[0].message.content.strip()

def extract_text_from_pdf(file):
    with pdfplumber.open(file) as pdf:
        return "\n".join(page.extract_text() for page in pdf.pages if page.extract_text())

def extract_text_from_docx(file):
    doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_details(text):
    name = text.split("\n")[0].strip() if text else ""
    email_match = re.search(r"[\w\.-]+@[\w\.-]+", text)
    phone_match = re.search(r"(\+\d{1,2}\s)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}", text)
    email = email_match.group() if email_match else ""
    phone = phone_match.group() if phone_match else ""
    skills = experience = ""
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if "skills" in line.lower():
            skills = "\n".join(lines[i+1:i+6])
        if "experience" in line.lower():
            experience = "\n".join(lines[i+1:i+10])
    return name, email, phone, skills, experience

def save_to_db(name, email, phone, skills, experience, filename):
    conn = sqlite3.connect(DB_FILE)
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO resumes (name, email, phone, skills, experience, filename, timestamp)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (name, email, phone, skills, experience, filename, datetime.now().isoformat()))
    conn.commit()
    conn.close()

def schedule_google_event(candidate_name, candidate_email, interview_date, interview_time, position_title, teams_link="https://teams.microsoft.com/l/meetup-join/abc123"):
    return teams_link or "https://teams.microsoft.com/l/meetup-join/abc123"

def generate_onboarding_doc(name, email, position, start_date, salary):
    doc = Document()
    doc.add_heading('Offer Letter', 0)
    doc.add_paragraph(f"Dear {name},")
    doc.add_paragraph(f"We are excited to offer you the position of {position}. Your start date will be {start_date}, with a starting salary of ${salary}.")
    doc.add_paragraph("Please let us know if you have any questions.")
    doc.add_paragraph("Sincerely,\nHR Team")

    for para in doc.paragraphs:
        for run in para.runs:
            run.font.size = Pt(11)

    filepath = f"onboarding_docs/{name.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
    os.makedirs("onboarding_docs", exist_ok=True)
    doc.save(filepath)

    conn = sqlite3.connect(DB_FILE)
    cur = conn.cursor()
    cur.execute("""
        INSERT INTO onboarding_logs (name, email, position, start_date, salary, filepath, timestamp)
        VALUES (?, ?, ?, ?, ?, ?, ?)
    """, (name, email, position, start_date, salary, filepath, datetime.now().isoformat()))
    conn.commit()
    conn.close()

    return filepath

# --- DB SETUP ---
def init_db():
    conn = sqlite3.connect(DB_FILE)
    cur = conn.cursor()

    cur.execute("""
    CREATE TABLE IF NOT EXISTS resumes (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT,
        email TEXT,
        phone TEXT,
        skills TEXT,
        experience TEXT,
        filename TEXT,
        timestamp TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS mira_logs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        question TEXT,
        answer TEXT,
        timestamp TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS onboarding_logs (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        name TEXT,
        email TEXT,
        position TEXT,
        start_date TEXT,
        salary TEXT,
        filepath TEXT,
        timestamp TEXT
    )
    """)

    cur.execute("""
    CREATE TABLE IF NOT EXISTS job_descriptions (
        id INTEGER PRIMARY KEY AUTOINCREMENT,
        content TEXT,
        timestamp TEXT
    )
    """)

    conn.commit()
    conn.close()

# --- UI ---
st.set_page_config(page_title="MIRA Assistant", layout="wide")
init_db()

def get_base64_image(image_path):
    with open(image_path, "rb") as f:
        return base64.b64encode(f.read()).decode()

mira_img_base64 = get_base64_image("mira.png")
st.markdown(f'''
    <style>
        .mira-header {{
            display: flex;
            flex-wrap: wrap;
            align-items: center;
            justify-content: flex-start;
            gap: 16px;
            text-align: left;
            margin-bottom: 20px;
        }}
        .mira-header img {{
            width: 80px;
            height: 80px;
            object-fit: cover;
            border-radius: 50%;
            border: 3px solid #d9a125;
        }}
        .mira-header h1 {{
            font-size: 1.8em;
            color: #a047fa;
            margin: 0;
        }}
    </style>
    <div class="mira-header">
        <img src="data:image/png;base64,{mira_img_base64}" />
        <h1>MIRA: Your AI Recruiting Assistant</h1>
    </div>
''', unsafe_allow_html=True)

TABS = [
    "🤖 Ask MIRA", 
    "📄 Resumes", 
    "🗕 Calendar & Onboarding", 
    "📁 Onboarding Docs", 
    "📂 Job Descriptions", 
    "📈 Upskilling & Coaching",
    "📚 MIRA Q&A Log"
]

tab1, tab2, tab3, tab4, tab5, tab6, tab7 = st.tabs(TABS)

# --- Render all tabs ---
def fetch_resumes(query=""):
    conn = sqlite3.connect(DB_FILE)
    cur = conn.cursor()
    if query:
        cur.execute("""
            SELECT * FROM resumes
            WHERE name LIKE ? OR email LIKE ? OR phone LIKE ?
            OR skills LIKE ? OR experience LIKE ?
        """, (f"%{query}%",)*5)
    else:
        cur.execute("SELECT * FROM resumes")
    rows = cur.fetchall()
    conn.close()
    return rows

def render_tabs(tab1, tab2, tab3, tab4, tab5, tab6, tab7):
    with tab1:
        st.subheader("🧠 How can I help?")
        user_input = st.text_input("Type your command here:")
        if user_input:
            try:
                response = ask_gpt(user_input)
                st.markdown(f"**MIRA says:** {response}")
                conn = sqlite3.connect(DB_FILE)
                conn.execute("INSERT INTO mira_logs (question, answer, timestamp) VALUES (?, ?, ?)", (user_input, response, datetime.now().isoformat()))
                conn.commit()
                conn.close()
            except Exception as e:
                st.error(f"Error: {e}")

    with tab2:
        st.subheader("📂 View Resumes")
        search = st.text_input("Search by name, email, skills, etc.")
        resumes = fetch_resumes(search)
        for r in resumes:
            st.markdown(f"**{r[1]}** | {r[2]} | {r[3]}")
            st.markdown(f"**Skills:** {r[4][:150]}...")
            st.markdown(f"**Experience:** {r[5][:200]}...")
            st.markdown("---")

    with tab3:
        st.subheader("📅 Interview Scheduling")
        with st.form("calendar_form"):
            name = st.text_input("Candidate Name")
            email = st.text_input("Candidate Email")
            position = st.text_input("Position Title")
            start_date = st.date_input("Start Date")
            salary = st.text_input("Offered Salary")
            interview_date = st.date_input("Interview Date")
            interview_time = st.time_input("Interview Time")
            teams_link = "https://teams.microsoft.com/l/meetup-join/abc123"

            if st.form_submit_button("📅 Schedule & Generate Onboarding"):
                try:
                    link = schedule_google_event(
                        name, email,
                        interview_date.strftime("%Y-%m-%d"),
                        interview_time.strftime("%H:%M"),
                        position,
                        teams_link
                    )
                    filepath = generate_onboarding_doc(name, email, position, start_date.strftime("%Y-%m-%d"), salary)
                    st.success(f"✅ Scheduled! Join link: {link}")
                    st.success(f"📄 Onboarding Doc: {filepath}")
                except Exception as e:
                    st.error(f"Error: {e}")

    with tab4:
        st.subheader("📁 Onboarding Docs")
        conn = sqlite3.connect(DB_FILE)
        rows = conn.execute("SELECT * FROM onboarding_logs ORDER BY timestamp DESC").fetchall()
        for r in rows:
            st.markdown(f"**{r[1]}** | {r[2]} | {r[3]} | ${r[5]} | Sent: {r[7]}")
            st.markdown(f"📄 File: {r[6]}")
            st.markdown("---")

    with tab5:
        st.subheader("📂 Job Descriptions")
        conn = sqlite3.connect(DB_FILE)
        jobs = conn.execute("SELECT * FROM job_descriptions ORDER BY timestamp DESC").fetchall()
        for jd in jobs:
            st.markdown(f"📝 **Generated:** {jd[2]}")
            st.code(jd[1])
            st.markdown("---")

    with tab6:
        st.subheader("📈 Upskilling & Coaching")
        st.info("Coming soon: AI-generated coaching tips and growth plans")

    with tab7:
        st.subheader("📚 Q&A Log")
        conn = sqlite3.connect(DB_FILE)
        logs = conn.execute("SELECT * FROM mira_logs ORDER BY timestamp DESC").fetchall()
        for row in logs:
            st.markdown(f"🕒 {row[3]}\n**Q:** {row[1]}\n**A:** {row[2]}")
            st.markdown("---")

render_tabs(tab1, tab2, tab3, tab4, tab5, tab6, tab7)